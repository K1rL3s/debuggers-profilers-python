import os
import re
import time
from typing import Optional, Set

import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)
FIGURE_FONT_SIZE = Pt(12)
CODE_FONT_NAME = "Courier New"
CODE_FONT_SIZE = Pt(12)
CODE_LINE_SPACING = 1.0
CODE_FIRST_LINE_INDENT = 0
PARAGRAPH_SPACE_BEFORE = Pt(0)
LINE_SPACING = 1.5
FIRST_LINE_INDENT = Cm(1.25)
TOP_MARGIN = Cm(2.0)
BOTTOM_MARGIN = Cm(2.0)
LEFT_MARGIN = Cm(3.0)
RIGHT_MARGIN = Cm(1.5)
IMAGE_WIDTH = Cm(16.5)
HYPERLINK_COLOR = RGBColor(0, 0, 255)
HEADING_COLOR = RGBColor(0, 0, 0)
TABLE_STYLE = "Table Grid"
MARKDOWN_EXTENSIONS = ["extra", "fenced_code", "tables"]
README_FILE = "README.md"
OUTPUT_FILE = f"Lesovoy_{int(time.time())}.docx"
MD_EXT = ".md"
PY_EXT = ".py"
PYX_EXT = ".pyx"
C_EXT = ".c"
H_EXT = ".h"
RS_EXT = ".rs"
TOML_EXT = ".toml"
CODE_BLOCK_SPACING = Pt(10)
HEADING_BASE_SIZE = 26
HEADING_SIZE_REDUCTION = 2
ERROR_PY_NOT_FOUND = "[Python file not found: {}]"
ERROR_MD_NOT_FOUND = "[Markdown file not found: {}]"
ERROR_IMAGE_NOT_FOUND = "[Image not found: {}]"
MAX_HEADING_LEVEL = 6
CODE_EXTENSIONS = (PY_EXT, PYX_EXT, C_EXT, H_EXT, RS_EXT, TOML_EXT)
IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg")

LISTING_COUNTER = 0
FIGURE_COUNTER = 0


def is_code_extension(filename: str) -> bool:
    return any(filename.endswith(ext) for ext in CODE_EXTENSIONS)


def is_image_extension(filename: str) -> bool:
    return any(filename.endswith(ext) for ext in IMAGE_EXTENSIONS)


def configure_document_style(document: Document) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE

    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = LINE_SPACING
    paragraph_format.first_line_indent = FIRST_LINE_INDENT
    paragraph_format.space_after = PARAGRAPH_SPACE_BEFORE
    paragraph_format.space_before = PARAGRAPH_SPACE_BEFORE

    for level in range(1, 6 + 1):
        heading_style = document.styles[f"Heading {level}"]
        heading_font = heading_style.font
        heading_font.name = FONT_NAME
        heading_font.size = Pt(HEADING_BASE_SIZE - level * HEADING_SIZE_REDUCTION)
        heading_font.color.rgb = HEADING_COLOR
        heading_font.bold = True if level <= 3 else False

        rPr = heading_style.element.rPr
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            heading_style.element.append(rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), FONT_NAME)
        rFonts.set(qn("w:hAnsi"), FONT_NAME)
        rFonts.set(qn("w:cs"), FONT_NAME)
        rFonts.set(qn("w:eastAsia"), FONT_NAME)

    if "Code" not in document.styles:
        code_style = document.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = document.styles["Code"]
    code_style.font.name = CODE_FONT_NAME
    code_style.font.size = CODE_FONT_SIZE
    code_style.paragraph_format.line_spacing = CODE_LINE_SPACING
    code_style.paragraph_format.space_after = CODE_BLOCK_SPACING
    code_style.paragraph_format.space_before = PARAGRAPH_SPACE_BEFORE
    code_style.paragraph_format.first_line_indent = CODE_FIRST_LINE_INDENT

    for section in document.sections:
        section.top_margin = TOP_MARGIN
        section.bottom_margin = BOTTOM_MARGIN
        section.left_margin = LEFT_MARGIN
        section.right_margin = RIGHT_MARGIN

    section = document.sections[0]
    footer = section.footer
    footer.paragraphs[0].text = "\t"
    run = footer.paragraphs[0].add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)
    section.different_first_page_header_footer = True
    document.sections[0].first_page_footer.paragraphs[0].text = ""


def add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    run = paragraph.add_run(text)
    run.font.color.rgb = HYPERLINK_COLOR
    run.font.underline = True

    run_xml = run._r
    paragraph._p.remove(run_xml)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.append(run_xml)
    paragraph._p.append(hyperlink)


def format_heading(heading, level: int) -> None:
    for run in heading.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(HEADING_BASE_SIZE - level * HEADING_SIZE_REDUCTION)
        run.font.color.rgb = HEADING_COLOR


def add_border_to_paragraph(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for border_name in ["w:top", "w:left", "w:bottom", "w:right"]:
        border = OxmlElement(border_name)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "000000")
        pBdr.append(border)
    pPr.append(pBdr)


def insert_code_block(
    document: Document,
    code_content: str,
    description: str | None = None,
) -> None:
    global LISTING_COUNTER
    LISTING_COUNTER += 1

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.first_line_indent = Cm(0)
    listing_text = f"Листинг {LISTING_COUNTER}"
    if description:
        listing_text += f" - {description}"
    run = caption.add_run(listing_text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE

    paragraph = document.add_paragraph(style="Code")
    run = paragraph.add_run(code_content.rstrip())
    add_border_to_paragraph(paragraph)


def insert_image(
    document: Document,
    image_path: str,
    description: str = "Изображение",
    width: Inches = IMAGE_WIDTH,
) -> None:
    global FIGURE_COUNTER
    if not os.path.exists(image_path):
        document.add_paragraph(ERROR_IMAGE_NOT_FOUND.format(image_path))
        return

    FIGURE_COUNTER += 1

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    run.add_picture(image_path, width=width)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    run = caption.add_run(f"Рисунок {FIGURE_COUNTER} - {description}")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FIGURE_FONT_SIZE


def insert_table(document: Document, table_html: str) -> None:
    soup = BeautifulSoup(table_html, "html.parser")
    table = soup.find("table")
    if not table:
        return

    rows = table.find_all("tr")
    if not rows:
        return

    num_rows = len(rows)
    num_cols = max(len(row.find_all(["td", "th"])) for row in rows)
    doc_table = document.add_table(rows=num_rows, cols=num_cols)
    doc_table.style = TABLE_STYLE

    for i, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        for j, cell in enumerate(cells):
            cell_text = cell.get_text(strip=True)
            doc_table.rows[i].cells[j].text = cell_text
            if cell.name == "th":
                for paragraph in doc_table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True


def handle_link(
    href: str,
    base_path: str,
    document: Document,
    paragraph,
    root_directory: str,
    level_increase: int,
    processed_files: Set[str],
    link_text: str,
    heading_level: Optional[int] = None,
) -> bool:
    if is_code_extension(href):
        py_path = os.path.normpath(os.path.join(base_path, href))
        if os.path.exists(py_path):
            with open(py_path, "r", encoding="utf-8") as f:
                code_content = f.read()
            insert_code_block(document, code_content, description=link_text)
        else:
            paragraph.add_run(ERROR_PY_NOT_FOUND.format(href))
        return True
    elif href.endswith(MD_EXT):
        md_path = os.path.normpath(os.path.join(base_path, href))
        if os.path.exists(md_path):
            heading_text = extract_h1_from_markdown(md_path, link_text)
            new_level = (
                heading_level if heading_level is not None else level_increase + 1
            )
            new_level = min(new_level, MAX_HEADING_LEVEL)
            heading = document.add_heading(heading_text, level=new_level)
            format_heading(heading, new_level)
            adjusted_level_increase = (
                new_level if heading_level is not None else level_increase
            )
            process_markdown(
                md_path,
                root_directory,
                document,
                level_increase=adjusted_level_increase,
                skip_h1=True,
                processed_files=processed_files,
            )
        else:
            paragraph.add_run(ERROR_MD_NOT_FOUND.format(href))
        return True
    elif is_image_extension(href):
        image_path = os.path.normpath(os.path.join(base_path, href))
        description = link_text
        insert_image(document, image_path, description)
        return True
    elif href:
        add_hyperlink(paragraph, href, link_text)
        return True
    return False


def extract_h1_from_markdown(
    file_path: str, fallback_text: Optional[str] = None
) -> str:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            markdown_content = f.read()
        if not markdown_content.strip():
            return fallback_text or os.path.basename(file_path)
        md = markdown.Markdown(extensions=MARKDOWN_EXTENSIONS)
        html_content = md.convert(markdown_content)
        soup = BeautifulSoup(html_content, "html.parser")
        h1 = soup.find("h1")
        return h1.get_text() if h1 else (fallback_text or os.path.basename(file_path))
    except Exception:
        return fallback_text or os.path.basename(file_path)


def adjust_headers(html_content: str, level_increase: int) -> str:
    soup = BeautifulSoup(html_content, "html.parser")
    for header in soup.find_all(re.compile("^h[1-9]$")):
        current_level = int(header.name[1])
        new_level = min(current_level + level_increase, MAX_HEADING_LEVEL)
        header.name = f"h{new_level}"
    return str(soup)


def process_list_element(
    list_element,
    document: Document,
    level_increase: int,
    processed_files: Set[str],
    list_type: str,
    nesting_level: int,
    markdown_file_path: str,
    root_directory: str,
) -> None:
    for li in list_element.find_all("li", recursive=False):
        paragraph = document.add_paragraph(
            style="List Bullet" if list_type == "bullet" else "List Number"
        )

        for child in li.children:
            if child.name == "a":
                href = child.get("href", "")
                link_text = child.get_text()
                handled = handle_link(
                    href,
                    os.path.dirname(markdown_file_path),
                    document,
                    paragraph,
                    root_directory,
                    level_increase,
                    processed_files,
                    link_text,
                )
                if handled:
                    continue
                paragraph.add_run(link_text)
            elif child.name == "strong":
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name == "em":
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == "code":
                run = paragraph.add_run(child.get_text())
                run.font.name = CODE_FONT_NAME
                run.font.size = CODE_FONT_SIZE
            elif child.name == "img":
                img_src = child.get("src", "")
                img_path = os.path.normpath(
                    os.path.join(os.path.dirname(markdown_file_path), img_src)
                )
                insert_image(
                    document, img_path, description=child.get("alt", "Изображение")
                )
            elif child.name == "ul":
                process_list_element(
                    child,
                    document,
                    level_increase,
                    processed_files,
                    list_type="bullet",
                    nesting_level=nesting_level + 1,
                    markdown_file_path=markdown_file_path,
                    root_directory=root_directory,
                )
            elif child.name == "ol":
                process_list_element(
                    child,
                    document,
                    level_increase,
                    processed_files,
                    list_type="number",
                    nesting_level=nesting_level + 1,
                    markdown_file_path=markdown_file_path,
                    root_directory=root_directory,
                )
            else:
                paragraph.add_run(str(child))


def process_markdown(
    markdown_file_path: str,
    root_directory: str,
    document: Document,
    level_increase: int = 0,
    skip_h1: bool = False,
    processed_files: Optional[Set[str]] = None,
) -> None:
    if processed_files is None:
        processed_files = set()

    if markdown_file_path in processed_files:
        return
    processed_files.add(markdown_file_path)

    with open(markdown_file_path, "r", encoding="utf-8") as f:
        markdown_content = f.read()

    md = markdown.Markdown(extensions=MARKDOWN_EXTENSIONS)
    html_content = md.convert(markdown_content)

    soup = BeautifulSoup(html_content, "html.parser")

    if skip_h1:
        for h1 in soup.find_all("h1"):
            h1.decompose()

    html_content = adjust_headers(str(soup), level_increase)
    soup = BeautifulSoup(html_content, "html.parser")

    for element in soup.children:
        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            heading_text = ""
            has_md_link = False

            for child in element.children:
                if child.name == "a" and child.get("href", "").endswith(MD_EXT):
                    href = child.get("href", "")
                    link_text = child.get_text()
                    paragraph = document.add_paragraph()
                    handled = handle_link(
                        href,
                        os.path.dirname(markdown_file_path),
                        document,
                        paragraph,
                        root_directory,
                        level_increase,
                        processed_files,
                        link_text,
                        heading_level=level,
                    )
                    if handled:
                        has_md_link = True
                        if paragraph.text:
                            document.paragraphs[-1]._element.getparent().remove(
                                document.paragraphs[-1]._element
                            )
                else:
                    heading_text += (
                        str(child) if child.name is None else child.get_text()
                    )

            if heading_text or not has_md_link:
                heading = document.add_heading(heading_text, level=level)
                format_heading(heading, level)

        elif element.name == "p":
            paragraph = document.add_paragraph()
            for child in element.children:
                if child.name == "a":
                    href = child.get("href", "")
                    link_text = child.get_text()
                    handled = handle_link(
                        href,
                        os.path.dirname(markdown_file_path),
                        document,
                        paragraph,
                        root_directory,
                        level_increase,
                        processed_files,
                        link_text,
                    )
                    if handled:
                        continue
                    paragraph.add_run(link_text)
                elif child.name == "strong":
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name == "em":
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                elif child.name == "code":
                    run = paragraph.add_run(child.get_text())
                    run.font.name = CODE_FONT_NAME
                    run.font.size = CODE_FONT_SIZE
                elif child.name == "img":
                    img_src = child.get("src", "")
                    img_path = os.path.normpath(
                        os.path.join(os.path.dirname(markdown_file_path), img_src)
                    )
                    insert_image(
                        document, img_path, description=child.get("alt", "Изображение")
                    )
                else:
                    paragraph.add_run(str(child))
        elif element.name == "ul":
            process_list_element(
                element,
                document,
                level_increase,
                processed_files,
                list_type="bullet",
                nesting_level=0,
                markdown_file_path=markdown_file_path,
                root_directory=root_directory,
            )
        elif element.name == "ol":
            process_list_element(
                element,
                document,
                level_increase,
                processed_files,
                list_type="number",
                nesting_level=0,
                markdown_file_path=markdown_file_path,
                root_directory=root_directory,
            )
        elif element.name == "table":
            insert_table(document, str(element))
        elif element.name == "pre":
            code = element.find("code")
            if code:
                insert_code_block(document, code.get_text())


def convert_markdown_to_docx(root_directory: str, output_docx: str) -> None:
    document = Document()
    configure_document_style(document)

    readme_path = os.path.join(root_directory, README_FILE)

    if not os.path.exists(readme_path):
        print(f"{README_FILE} not found!")
        return

    with open(readme_path, "r", encoding="utf-8") as f:
        readme_content = f.read()

    md = markdown.Markdown(extensions=MARKDOWN_EXTENSIONS)
    html_content = md.convert(readme_content)
    soup = BeautifulSoup(html_content, "html.parser")

    processed_files = set()
    current_heading_level = 1

    for element in soup.children:
        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            heading_text = ""
            has_md_link = False

            for child in element.children:
                if child.name == "a" and child.get("href", "").endswith(MD_EXT):
                    href = child.get("href", "")
                    link_text = child.get_text()
                    paragraph = document.add_paragraph()
                    handled = handle_link(
                        href,
                        root_directory,
                        document,
                        paragraph,
                        root_directory,
                        current_heading_level,
                        processed_files,
                        link_text,
                        heading_level=level,
                    )
                    if handled:
                        has_md_link = True
                        if paragraph.text:
                            document.paragraphs[-1]._element.getparent().remove(
                                document.paragraphs[-1]._element
                            )
                else:
                    heading_text += (
                        str(child) if child.name is None else child.get_text()
                    )

            if heading_text or not has_md_link:
                heading = document.add_heading(heading_text, level=level)
                format_heading(heading, level)

        elif element.name == "p":
            paragraph = document.add_paragraph()
            for child in element.children:
                if child.name == "a":
                    href = child.get("href", "")
                    link_text = child.get_text()
                    handled = handle_link(
                        href,
                        root_directory,
                        document,
                        paragraph,
                        root_directory,
                        current_heading_level,
                        processed_files,
                        link_text,
                    )
                    if handled:
                        continue
                    paragraph.add_run(link_text)
                elif child.name == "strong":
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name == "em":
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                elif child.name == "code":
                    run = paragraph.add_run(child.get_text())
                    run.font.name = CODE_FONT_NAME
                    run.font.size = CODE_FONT_SIZE
                elif child.name == "img":
                    img_src = child.get("src", "")
                    img_path = os.path.normpath(os.path.join(root_directory, img_src))
                    insert_image(
                        document, img_path, description=child.get("alt", "Изображение")
                    )
                else:
                    paragraph.add_run(str(child))
        elif element.name == "ul":
            process_list_element(
                element,
                document,
                current_heading_level,
                processed_files,
                list_type="bullet",
                nesting_level=0,
                markdown_file_path=readme_path,
                root_directory=root_directory,
            )
        elif element.name == "ol":
            process_list_element(
                element,
                document,
                current_heading_level,
                processed_files,
                list_type="number",
                nesting_level=0,
                markdown_file_path=readme_path,
                root_directory=root_directory,
            )
        elif element.name == "table":
            insert_table(document, str(element))
        elif element.name == "pre":
            code = element.find("code")
            if code:
                insert_code_block(document, code.get_text())

    document.save(output_docx)
    print(f"Document saved as {output_docx}")


if __name__ == "__main__":
    root_directory = "."
    output_docx_file = OUTPUT_FILE
    convert_markdown_to_docx(root_directory, output_docx_file)
